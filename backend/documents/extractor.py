"""Text extraction from document files — cahier des charges §13, §9.

This is the missing front end of the documentary workflow. Chunking,
embeddings, indexing and semantic search already existed (Echo +
backend/memory/semantic.py), but `memory_index` only ever accepted text
that had *already* been extracted — so "Import → OCR → Découpage →
Embeddings" (§9) stopped at its first step for anything that wasn't
already a string.

Deliberately knows nothing about Aegis, Echo, or HTTP: it turns bytes
into text and nothing else, which is what makes it testable without a
running Ollama, ChromaDB, or security engine. The Aegis-gated read and
the indexing are composed on top of it (see backend/mcp_server/server.py's
documents_index).

Dependency policy (backend/requirements.txt: "phase by phase"): the
plain-text family needs nothing at all, PDF and DOCX need one pure-Python
library each, and both are imported *inside* their extractor rather than
at module level. A missing library therefore degrades to a clear,
actionable error on that one format instead of breaking the import of
every module that touches documents.

Images are intentionally not handled here — see _extract_image.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

# Extensions handled as plain text. Code, config and data files all land
# here: for RAG purposes their bytes *are* their content, so there is
# nothing to "extract" beyond decoding them.
_TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c", ".h", ".cpp", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".ps1", ".sql", ".r",
    ".html", ".htm", ".css", ".scss", ".xml", ".csv", ".tsv", ".log",
}

_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"}

# Tried in order. utf-8-sig comes FIRST, not after utf-8: plain utf-8
# decodes a BOM quite happily into a leading U+FEFF, so listing it first
# means utf-8-sig never gets a turn and the BOM survives into the first
# chunk. utf-8-sig strips the BOM when present and is identical to utf-8
# when absent, so it is strictly the better first attempt. cp1252 catches
# Windows-authored files, the common case on this machine. The final pass
# cannot fail — a few replacement characters beat refusing to index.
_ENCODINGS = ("utf-8-sig", "cp1252", "latin-1")


class UnsupportedDocumentError(ValueError):
    """Raised for a format this extractor deliberately does not handle."""


class MissingExtractorDependencyError(RuntimeError):
    """Raised when the format is supported but its library isn't installed.

    Separate from UnsupportedDocumentError on purpose: one means "never
    going to work here", the other means "run one pip install". Callers
    (and users) should be able to tell those apart.
    """


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    format: str
    # Number of source units the text came from — PDF pages, DOCX
    # paragraphs, 1 for a plain file. Surfaced so a caller can tell an
    # empty extraction from a genuinely empty document.
    units: int = 1
    warnings: tuple[str, ...] = field(default_factory=tuple)


def supported_suffixes() -> set[str]:
    """Everything extract_text() will attempt, images excluded."""
    return _TEXT_SUFFIXES | {".pdf", ".docx"}


def _normalize_newlines(text: str) -> str:
    """CRLF/CR → LF. Chunk boundaries are computed on character counts
    (semantic.py's chunk_text), so the same document authored on Windows
    and on Linux would otherwise chunk differently and embed differently
    — a difference with no meaning that RAG results would inherit."""
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _decode(data: bytes) -> tuple[str, str | None]:
    for encoding in _ENCODINGS:
        try:
            return _normalize_newlines(data.decode(encoding)), None
        except UnicodeDecodeError:
            continue
    # latin-1 above decodes any byte sequence, so this is unreachable in
    # practice; kept so a future change to _ENCODINGS can't silently
    # produce an unbound variable.
    return (
        _normalize_newlines(data.decode("utf-8", errors="replace")),
        "decoded with replacement characters",
    )


def _extract_text_file(data: bytes, suffix: str) -> ExtractedDocument:
    text, warning = _decode(data)
    return ExtractedDocument(
        text=text,
        format=suffix.lstrip("."),
        units=1,
        warnings=(warning,) if warning else (),
    )


def _extract_pdf(data: bytes) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - exercised via monkeypatch
        raise MissingExtractorDependencyError(
            "PDF extraction needs the 'pypdf' package: pip install pypdf"
        ) from exc

    import io

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    warnings: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - one bad page shouldn't lose the rest
            pages.append("")
            warnings.append(f"page {number} could not be extracted: {exc}")

    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        # A scanned PDF is images in a PDF wrapper: pypdf finds no text
        # layer and returns empty strings. Saying so is far more useful
        # than indexing an empty document and wondering later why the
        # search never matches it.
        warnings.append(
            "no text layer found — this looks like a scanned PDF; "
            "extract its pages as images and use analyze_image instead"
        )
    return ExtractedDocument(
        text=text, format="pdf", units=len(reader.pages), warnings=tuple(warnings)
    )


def _extract_docx(data: bytes) -> ExtractedDocument:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - exercised via monkeypatch
        raise MissingExtractorDependencyError(
            "DOCX extraction needs the 'python-docx' package: pip install python-docx"
        ) from exc

    import io

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry real content in specs and reports; python-docx keeps
    # them out of .paragraphs, so they'd be silently dropped otherwise.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractedDocument(text="\n".join(parts), format="docx", units=len(parts))


def _extract_image(suffix: str) -> ExtractedDocument:
    raise UnsupportedDocumentError(
        f"{suffix} is an image. This project has a vision model already "
        "(Hermes Eyes / gemma4) — use the analyze_image tool, which reads "
        "diagrams and screenshots, not just glyphs. Adding an OCR stack "
        "(tesseract) would be a heavier dependency for a narrower result."
    )


def extract_text(path: str | Path, *, data: bytes | None = None) -> ExtractedDocument:
    """Extract indexable text from a document.

    `data` lets a caller supply bytes it has already read — the MCP tool
    does exactly that, because its read has to go through Aegis rather
    than touching the filesystem here.
    """
    target = Path(path)
    suffix = target.suffix.lower()

    if data is None:
        if not target.exists():
            raise FileNotFoundError(f"No such file: {target}")
        data = target.read_bytes()

    if suffix in _IMAGE_SUFFIXES:
        return _extract_image(suffix)
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in _TEXT_SUFFIXES:
        return _extract_text_file(data, suffix)
    if suffix == ".doc":
        raise UnsupportedDocumentError(
            ".doc is the pre-2007 binary Word format, which python-docx "
            "cannot read. Convert it to .docx first."
        )

    raise UnsupportedDocumentError(
        f"Unsupported extension {suffix!r}. Supported: "
        f"{', '.join(sorted(supported_suffixes()))}"
    )
